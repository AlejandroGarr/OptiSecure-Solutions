from docx import Document
from docx.shared import Pt, RGBColor

doc = Document(r'c:\Users\aleja\Desktop\TFG\OptiSecure-Solutions\TFG_PABLO_GARCÍA_RAMÍREZ.docx')

# Find where content chapters start
for i, p in enumerate(doc.paragraphs[55:150]):
    sname = p.style.name if p.style else 'None'
    txt = p.text[:120].replace(chr(10), ' ')
    if txt.strip():
        print(f'[{i+55:3d}] [{sname:30s}] {txt}')

for style_name in ['Heading 1', 'Heading 2', 'Heading 4', 'Body Text', 'Title', 'List Paragraph', 'Normal']:
    print(f'\n=== {style_name} ===')
    try:
        s = doc.styles[style_name]
        f = s.font
        color_str = str(f.color.rgb) if f.color and f.color.rgb else 'inherited'
        print(f'  Name: {f.name}, Size: {f.size}, Bold: {f.bold}, Color: {color_str}')
        pf = s.paragraph_format
        print(f'  Alignment: {pf.alignment}, SpaceBefore: {pf.space_before}, SpaceAfter: {pf.space_after}')
        print(f'  LineSpacing: {pf.line_spacing}')
    except Exception as e:
        print(f'  Error: {e}')

# Check for tables
print(f'\n=== TABLES: {len(doc.tables)} ===')
for i, t in enumerate(doc.tables[:3]):
    print(f'Table {i}: {len(t.rows)} rows x {len(t.columns)} cols')
    for j, row in enumerate(t.rows[:2]):
        cells = [c.text[:30] for c in row.cells]
        print(f'  Row {j}: {cells}')

# Check sections (page setup)
print(f'\n=== SECTIONS ===')
for i, sec in enumerate(doc.sections):
    print(f'Section {i}: width={sec.page_width}, height={sec.page_height}')
    print(f'  margins: top={sec.top_margin}, bottom={sec.bottom_margin}, left={sec.left_margin}, right={sec.right_margin}')
